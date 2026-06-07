from docx import Document
from src.models import MeetingOutput
from pathlib import Path
import logging

from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def generate_document(meeting: MeetingOutput, output_path: Path) -> None:
    """
    Generates a formatted Word document from the MeetingOutput object.

    Creates a professional-looking report with sections for summary, key points,
    action items, risks, and decisions.

    Args:
        meeting: Structured meeting data to format into the document
        output_path: Path where the document should be saved
    """
    logger.info("Generating document for meeting")
    doc = Document()

    # Main title - level=0 is biggest, only used once
    title = doc.add_heading("Meeting Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(meeting.summary)

    # Key Points section
    if meeting.key_points:
        doc.add_heading("Key Points", level=1)
        for point in meeting.key_points:
            doc.add_paragraph(point, style="List Bullet")

    # Action Items table
    if meeting.action_items:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        
        header = table.rows[0].cells
        header[0].text = "Task"
        header[1].text = "Owner"
        header[2].text = "Deadline"
        header[3].text = "Priority"
        
        for item in meeting.action_items:
            row = table.add_row().cells
            row[0].text = item.task
            row[1].text = item.owner
            row[2].text = item.deadline
            row[3].text = item.priority

    # Risks table
    if meeting.risks:
        doc.add_heading("Risks", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        header = table.rows[0].cells
        header[0].text = "Description"
        header[1].text = "Severity"
        
        for risk in meeting.risks:
            row = table.add_row().cells
            row[0].text = risk.description
            row[1].text = risk.severity

    # Decisions section
    if meeting.decisions:
        doc.add_heading("Decisions", level=1)
        for decision in meeting.decisions:
            doc.add_paragraph(decision, style="List Bullet")

    doc.save(output_path)
    logger.info(f"Document generated successfully: {output_path}")

